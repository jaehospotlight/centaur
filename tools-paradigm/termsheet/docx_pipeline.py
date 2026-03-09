"""Fill Paradigm DOCX term sheet templates with tracked-redline support."""

from __future__ import annotations

import hashlib
import io
import logging
import os
import posixpath
import re
import tempfile
import zipfile
from collections.abc import Iterable
from copy import deepcopy
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any

import httpx
from docx import Document
from docx.text.run import Run
from docx_revisions import RevisionDocument  # pyright: ignore[reportMissingImports]
from lxml import etree

from .models import BoardRights, InstrumentType, TermSheet

log = logging.getLogger(__name__)

TEMPLATE_DIR = Path(__file__).resolve().parent / "templates"
TEMPLATE_PATH = TEMPLATE_DIR / "paradigm_term_sheet.docx"
PREPROCESSED_TEMPLATE_PATH = TEMPLATE_DIR / "paradigm_term_sheet_preprocessed.docx"

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NS = {"w": _W_NS}
_REDLINE_AUTHOR = "Paradigm Legal"


@dataclass
class GeneratedPackage:
    clean_docx: bytes
    redline_docx: bytes | None
    clean_pdf: bytes | None
    redline_pdf: bytes | None
    replacements: dict[str, str]
    fidelity_report: dict[str, Any]


def format_money(amount: float) -> str:
    """Short-form: $50M, $7.5M, $500K."""
    if amount >= 1_000_000_000:
        return f"${amount / 1_000_000_000:g}B"
    if amount >= 1_000_000:
        return f"${amount / 1_000_000:g}M"
    if amount >= 1_000:
        return f"${amount / 1_000:g}K"
    return f"${amount:,.0f}"


def format_money_full(amount: float) -> str:
    """Full-form: $75,000  $100,000."""
    return f"${amount:,.0f}"


def _iter_all_paragraphs(doc: Document) -> Iterable:
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for paragraph in section.footer.paragraphs:
            yield paragraph


def _run_boundaries(paragraph) -> list[tuple[int, int]]:
    boundaries: list[tuple[int, int]] = []
    cursor = 0
    for run in paragraph.runs:
        text = run.text or ""
        next_cursor = cursor + len(text)
        boundaries.append((cursor, next_cursor))
        cursor = next_cursor
    return boundaries


def _split_run(paragraph, run_index: int, local_offset: int) -> None:
    run = paragraph.runs[run_index]
    text = run.text or ""
    if local_offset <= 0 or local_offset >= len(text):
        return
    run.text = text[:local_offset]
    copied = deepcopy(run._r)
    run._r.addnext(copied)
    new_run = Run(copied, run._parent)
    new_run.text = text[local_offset:]


def _split_at_absolute_index(paragraph, absolute_index: int) -> None:
    if absolute_index <= 0:
        return
    boundaries = _run_boundaries(paragraph)
    for idx, (start, end) in enumerate(boundaries):
        if start < absolute_index < end:
            _split_run(paragraph, idx, absolute_index - start)
            return


def _isolate_run(paragraph, start: int, end: int):
    _split_at_absolute_index(paragraph, end)
    _split_at_absolute_index(paragraph, start)
    boundaries = _run_boundaries(paragraph)
    for idx, (run_start, _run_end) in enumerate(boundaries):
        if run_start == start:
            return paragraph.runs[idx]
    return None


def _replace_span(paragraph, start: int, end: int, replacement: str) -> bool:
    _isolate_run(paragraph, start, end)
    boundaries = _run_boundaries(paragraph)
    covered: list[int] = []
    for idx, (rs, re_) in enumerate(boundaries):
        if rs >= start and re_ <= end:
            covered.append(idx)
    if not covered:
        return False
    paragraph.runs[covered[0]].text = replacement
    for idx in covered[1:]:
        paragraph.runs[idx].text = ""
    return True


def _replace_regex_once(paragraph, pattern: str, replacement: str, flags: int = 0) -> bool:
    match = re.search(pattern, paragraph.text, flags)
    if match is None:
        return False
    return _replace_span(paragraph, match.start(), match.end(), replacement)


def _replace_regex_all(paragraph, pattern: str, replacement: str, flags: int = 0) -> int:
    replaced = 0
    while True:
        before = paragraph.text
        if not _replace_regex_once(paragraph, pattern, replacement, flags):
            break
        if paragraph.text == before:
            break
        replaced += 1
    return replaced


def _strip_proof_errors(docx_bytes: bytes) -> bytes:
    src = io.BytesIO(docx_bytes)
    out = io.BytesIO()
    with (
        zipfile.ZipFile(src, "r") as zin,
        zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout,
    ):
        for info in zin.infolist():
            data = zin.read(info.filename)
            if info.filename.startswith("word/") and info.filename.endswith(".xml"):
                try:
                    tree = etree.fromstring(data)
                    for node in tree.xpath("//w:proofErr", namespaces=_NS):
                        parent = node.getparent()
                        if parent is not None:
                            parent.remove(node)
                    data = etree.tostring(tree, encoding="utf-8", xml_declaration=True)
                except etree.XMLSyntaxError:
                    pass
            zout.writestr(info, data)
    return out.getvalue()


def _same_run_format(left, right) -> bool:
    if left._r.rPr is None and right._r.rPr is None:
        return True
    if left._r.rPr is None or right._r.rPr is None:
        return False
    return etree.tostring(left._r.rPr) == etree.tostring(right._r.rPr)


def _merge_identical_runs(paragraph) -> None:
    if len(paragraph.runs) < 2:
        return
    idx = 0
    while idx < len(paragraph.runs) - 1:
        current = paragraph.runs[idx]
        nxt = paragraph.runs[idx + 1]
        if _same_run_format(current, nxt):
            current.text = f"{current.text}{nxt.text}"
            nxt._r.getparent().remove(nxt._r)
        else:
            idx += 1


def preprocess_template(
    template_path: str | Path = TEMPLATE_PATH,
    output_path: str | Path = PREPROCESSED_TEMPLATE_PATH,
) -> str:
    template_path = Path(template_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    raw = template_path.read_bytes()
    stripped = _strip_proof_errors(raw)
    doc = Document(io.BytesIO(stripped))
    for paragraph in _iter_all_paragraphs(doc):
        _merge_identical_runs(paragraph)
    doc.save(str(output_path))
    return str(output_path)


def _cached_preprocessed_template_path(template_path: Path) -> Path:
    digest = hashlib.sha256(template_path.read_bytes()).hexdigest()[:16]
    cache_dir = Path(tempfile.gettempdir()) / "ai_v2_termsheet_cache"
    cache_dir.mkdir(parents=True, exist_ok=True)
    return cache_dir / f"{template_path.stem}-{digest}-preprocessed.docx"


def _ensure_preprocessed_template(template_path: Path) -> Path:
    preprocessed_path = _cached_preprocessed_template_path(template_path)
    if not preprocessed_path.exists():
        preprocess_template(template_path=template_path, output_path=preprocessed_path)
    return preprocessed_path


def _build_replacements(ts: TermSheet) -> dict[str, str]:
    valuation = ts.effective_valuation
    ownership = ts.ownership_percent
    series = ts.effective_series
    return {
        "company_upper": ts.company_name.upper(),
        "series": series,
        "investment": format_money(ts.investment_amount),
        "valuation": format_money(valuation) if valuation else "",
        "option_pool": f"{ts.option_pool_percent:g}",
        "ownership": f"{ownership:.1f}" if ownership is not None else "",
        "debt_threshold": format_money(ts.debt_threshold),
        "token_floor": f"{ts.token_rights.token_floor_percent:g}",
        "ipo_threshold": format_money(ts.ipo_threshold),
        "legal_fee_cap": format_money_full(ts.legal_fee_cap),
        "nvca_year": str(ts.nvca_year),
        "exclusivity_days": str(ts.exclusivity_days),
        "founder_carveout": f"{ts.founder_carveout_percent:g}",
    }


def _apply_board_language(paragraph, board_rights: BoardRights) -> None:
    """Select director/observer clauses from template brackets."""
    text = paragraph.text
    if "preferred stock voting thresholds" not in text.lower():
        return
    if "[" not in text or "]" not in text:
        return
    match = re.match(
        r"^\s*\[(.*?)\]\s*\[(.*?)\]\s*(Preferred Stock voting thresholds.*)$",
        text,
    )
    if not match:
        return
    director_clause = match.group(1).strip()
    observer_clause = match.group(2).strip()
    tail = match.group(3).strip()

    chosen: list[str] = []
    if board_rights in {BoardRights.SEAT, BoardRights.SEAT_AND_OBSERVER}:
        chosen.append(director_clause.rstrip(";.") + ".")
    if board_rights in {BoardRights.OBSERVER, BoardRights.SEAT_AND_OBSERVER}:
        obs = observer_clause
        if board_rights == BoardRights.OBSERVER:
            obs = re.sub(r"^In addition,\s*", "", obs, flags=re.IGNORECASE)
            obs = re.sub(
                r",?\s*if Paradigm has not designated its director,?\s*",
                " ",
                obs,
                flags=re.IGNORECASE,
            )
        chosen.append(obs.rstrip(";.") + ".")

    result = " ".join([*chosen, tail]).strip()
    _replace_span(paragraph, 0, len(text), result)


def _apply_co_investor_language(paragraph, include: bool, custom_text: str | None = None) -> None:
    text = paragraph.text
    match = re.search(
        r"\[Other investors.*?post-money valuation\]\.?",
        text,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if match is None:
        return
    replacement = ""
    if include:
        replacement = (
            custom_text
            or "Other investors mutually acceptable to Paradigm and the Company may "
            "invest additional amounts, which shall not affect the post-money valuation."
        )
    _replace_span(paragraph, match.start(), match.end(), replacement)


def _apply_ipo_threshold(paragraph, threshold_str: str) -> None:
    """Replace IPO threshold in Securities section."""
    _replace_regex_once(
        paragraph,
        r"net proceeds greater than \$[\d,.]+[MBK]?",
        f"net proceeds greater than {threshold_str}",
    )


def _apply_legal_fee_cap(paragraph, fee_cap_str: str) -> None:
    _replace_regex_once(
        paragraph,
        r"counsel up to \$[\d,]+",
        f"counsel up to {fee_cap_str}",
    )


def _apply_nvca_year(paragraph, year: str) -> None:
    _replace_regex_all(paragraph, r"\d{4} NVCA forms", f"{year} NVCA forms")


def _apply_exclusivity_days(paragraph, days: str) -> None:
    _replace_regex_once(
        paragraph,
        r"period of \d+ days",
        f"period of {days} days",
    )


def _apply_founder_carveout(paragraph, pct: str) -> None:
    _replace_regex_once(
        paragraph,
        r"up to \d+% of the stock initially",
        f"up to {pct}% of the stock initially",
    )


def _apply_pro_rata_toggle(paragraph, enabled: bool) -> None:
    if enabled:
        return
    _replace_regex_once(
        paragraph,
        r"including information rights and pro rata rights \(including overallotment\) for Major Investors \(which shall only include Paradigm\),\s*",
        "including information rights, ",
    )


def _apply_seed_language(paragraph, is_seed: bool) -> None:
    """For seed deals, remove 'together with other series of Preferred Stock, '."""
    if not is_seed:
        return
    _replace_regex_all(
        paragraph,
        r"together with other series of Preferred Stock, ",
        "",
    )


def _apply_row_override(paragraph, anchor: str, replacement: str | None) -> None:
    if replacement and anchor.lower() in paragraph.text.lower():
        _replace_span(paragraph, 0, len(paragraph.text), replacement.strip())


def _apply_protective_v_override(paragraph, clause_v_text: str | None) -> None:
    if not clause_v_text:
        return
    if "(v)" not in paragraph.text or "related party transactions" not in paragraph.text.lower():
        return
    clause = clause_v_text.strip().rstrip(".")
    _replace_regex_once(
        paragraph,
        r"\(v\)\s+any interested or related party transactions.*$",
        f"(v) {clause}.",
    )


def _remove_token_rights_rows(doc: Document) -> None:
    for table in doc.tables:
        rows_to_remove = []
        for row in table.rows:
            label = row.cells[0].text.strip().lower() if row.cells else ""
            if "token rights" in label:
                rows_to_remove.append(row)
        for row in rows_to_remove:
            table._tbl.remove(row._tr)


def _word_xml_blob(docx_bytes: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as archive:
        parts = [
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        ]
    return "\n".join(parts)


def _assert_no_unresolved_markers(docx_bytes: bytes) -> None:
    xml = _word_xml_blob(docx_bytes)
    unresolved_patterns = (
        r"\$\[__\]M",
        r"\[__\]\s*%",
        r"\[COMPANY\]",
        r"SERIES\s*\[\s*[_ ]\s*\]",
        r"Series\s*\[__\]",
        r"\[Token Floor usually 50%\]",
        r"\[Other investors.*?post-money valuation\]",
        r"\{\{[^{}]+\}\}",
    )
    hits: list[str] = []
    for pattern in unresolved_patterns:
        if re.search(pattern, xml, flags=re.IGNORECASE | re.DOTALL):
            hits.append(pattern)
    if hits:
        raise ValueError(f"Unresolved template markers remain: {', '.join(hits)}")


def _assert_template_structure_preserved(template_docx: bytes, generated_docx: bytes) -> None:
    with zipfile.ZipFile(io.BytesIO(template_docx), "r") as src, zipfile.ZipFile(
        io.BytesIO(generated_docx), "r"
    ) as out:
        src_names = set(src.namelist())
        out_names = set(out.namelist())

    src_headers = {name for name in src_names if name.startswith("word/header") and name.endswith(".xml")}
    out_headers = {name for name in out_names if name.startswith("word/header") and name.endswith(".xml")}
    if src_headers != out_headers:
        raise ValueError("Header XML parts changed unexpectedly from template")

    src_header_rels = {
        name for name in src_names if name.startswith("word/_rels/header") and name.endswith(".rels")
    }
    out_header_rels = {
        name for name in out_names if name.startswith("word/_rels/header") and name.endswith(".rels")
    }
    if src_header_rels != out_header_rels:
        raise ValueError("Header relationship parts changed unexpectedly from template")

    src_media = {name for name in src_names if name.startswith("word/media/")}
    out_media = {name for name in out_names if name.startswith("word/media/")}
    if not src_media.issubset(out_media):
        raise ValueError("Template media assets missing from generated document")


def _docx_parts(docx_bytes: bytes) -> dict[str, bytes]:
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as archive:
        return {name: archive.read(name) for name in archive.namelist()}


def _sha256_bytes(value: bytes) -> str:
    return hashlib.sha256(value).hexdigest()


def _normalize_docx_path(path: str) -> str:
    normalized = posixpath.normpath(path)
    return normalized.lstrip("./")


def _resolve_relationship_target(rel_path: str, target: str) -> str:
    rel_dir = posixpath.dirname(rel_path)
    source_name = posixpath.basename(rel_path)
    owner_dir = posixpath.dirname(rel_dir)
    source_part = posixpath.join(owner_dir, source_name[: -len(".rels")])
    source_dir = posixpath.dirname(source_part)
    return _normalize_docx_path(posixpath.join(source_dir, target))


def _header_banner_targets(parts: dict[str, bytes]) -> list[str]:
    targets: set[str] = set()
    for name, content in parts.items():
        if not name.startswith("word/_rels/header") or not name.endswith(".rels"):
            continue
        try:
            root = etree.fromstring(content)
        except etree.XMLSyntaxError:
            continue
        for rel in root.findall("{http://schemas.openxmlformats.org/package/2006/relationships}Relationship"):
            rel_type = rel.attrib.get("Type", "")
            rel_target = rel.attrib.get("Target", "")
            if rel_type.endswith("/image") and rel_target:
                targets.add(_resolve_relationship_target(name, rel_target))
    return sorted(targets)


def _font_names(parts: dict[str, bytes]) -> list[str]:
    font_table = parts.get("word/fontTable.xml")
    if not font_table:
        return []
    try:
        root = etree.fromstring(font_table)
    except etree.XMLSyntaxError:
        return []
    names: list[str] = []
    for node in root.findall(".//w:font", namespaces=_NS):
        font_name = node.attrib.get(f"{{{_W_NS}}}name")
        if font_name:
            names.append(font_name)
    return sorted(set(names))


def _style_ids(parts: dict[str, bytes]) -> list[str]:
    styles_part = parts.get("word/styles.xml")
    if not styles_part:
        return []
    try:
        root = etree.fromstring(styles_part)
    except etree.XMLSyntaxError:
        return []
    values: list[str] = []
    for node in root.findall(".//w:style", namespaces=_NS):
        style_id = node.attrib.get(f"{{{_W_NS}}}styleId")
        if style_id:
            values.append(style_id)
    return sorted(set(values))


def _immutable_template_part_names(template_parts: dict[str, bytes]) -> set[str]:
    names = {
        part
        for part in (
            "word/styles.xml",
            "word/fontTable.xml",
            "word/theme/theme1.xml",
            "word/settings.xml",
        )
        if part in template_parts
    }
    names.update(
        name for name in template_parts if name.startswith("word/header") and name.endswith(".xml")
    )
    names.update(
        name
        for name in template_parts
        if name.startswith("word/_rels/header") and name.endswith(".rels")
    )
    names.update(name for name in template_parts if name.startswith("word/media/"))
    return names


def _restore_template_immutable_parts(template_docx: bytes, generated_docx: bytes) -> bytes:
    template_parts = _docx_parts(template_docx)
    immutable_parts = _immutable_template_part_names(template_parts)
    if not immutable_parts:
        return generated_docx

    generated_buffer = io.BytesIO(generated_docx)
    output_buffer = io.BytesIO()
    with zipfile.ZipFile(generated_buffer, "r") as source_zip, zipfile.ZipFile(
        output_buffer, "w", zipfile.ZIP_DEFLATED
    ) as output_zip:
        generated_names = set(source_zip.namelist())
        for info in source_zip.infolist():
            replacement = template_parts.get(info.filename) if info.filename in immutable_parts else None
            output_zip.writestr(info, replacement if replacement is not None else source_zip.read(info.filename))
        for name in sorted(immutable_parts):
            if name not in generated_names:
                output_zip.writestr(name, template_parts[name])
    return output_buffer.getvalue()


def _build_template_fidelity_report(template_docx: bytes, generated_docx: bytes) -> dict[str, Any]:
    template_parts = _docx_parts(template_docx)
    generated_parts = _docx_parts(generated_docx)

    protected_parts = [
        "word/styles.xml",
        "word/fontTable.xml",
        "word/theme/theme1.xml",
        "word/settings.xml",
    ]
    protected_hashes_match = True
    protected_parts_present = True
    protected_part_hashes: dict[str, dict[str, str]] = {}
    for part in protected_parts:
        template_blob = template_parts.get(part)
        generated_blob = generated_parts.get(part)
        if template_blob is None:
            continue
        if generated_blob is None:
            protected_hashes_match = False
            protected_parts_present = False
            continue
        template_hash = _sha256_bytes(template_blob)
        generated_hash = _sha256_bytes(generated_blob)
        protected_part_hashes[part] = {
            "template_sha256": template_hash,
            "generated_sha256": generated_hash,
        }
        if template_hash != generated_hash:
            protected_hashes_match = False

    template_header_parts = sorted(
        name for name in template_parts if name.startswith("word/header") and name.endswith(".xml")
    )
    generated_header_parts = sorted(
        name for name in generated_parts if name.startswith("word/header") and name.endswith(".xml")
    )
    header_parts_present = template_header_parts == generated_header_parts
    header_parts_unchanged = all(
        _sha256_bytes(template_parts[name]) == _sha256_bytes(generated_parts.get(name, b""))
        for name in template_header_parts
    )

    template_header_rel_parts = sorted(
        name for name in template_parts if name.startswith("word/_rels/header") and name.endswith(".rels")
    )
    generated_header_rel_parts = sorted(
        name for name in generated_parts if name.startswith("word/_rels/header") and name.endswith(".rels")
    )
    header_rel_parts_present = template_header_rel_parts == generated_header_rel_parts
    header_rel_parts_unchanged = all(
        _sha256_bytes(template_parts[name]) == _sha256_bytes(generated_parts.get(name, b""))
        for name in template_header_rel_parts
    )

    banner_targets = _header_banner_targets(template_parts)
    banner_integrity = all(
        target in generated_parts
        and _sha256_bytes(template_parts[target]) == _sha256_bytes(generated_parts[target])
        for target in banner_targets
    )

    template_fonts = _font_names(template_parts)
    generated_fonts = _font_names(generated_parts)
    fonts_unchanged = template_fonts == generated_fonts

    template_style_ids = _style_ids(template_parts)
    generated_style_ids = _style_ids(generated_parts)
    style_ids_unchanged = template_style_ids == generated_style_ids

    media_parts = sorted(name for name in template_parts if name.startswith("word/media/"))
    media_unchanged = all(
        name in generated_parts
        and _sha256_bytes(template_parts[name]) == _sha256_bytes(generated_parts[name])
        for name in media_parts
    )

    passed = all(
        (
            protected_parts_present,
            protected_hashes_match,
            header_parts_present,
            header_parts_unchanged,
            header_rel_parts_present,
            header_rel_parts_unchanged,
            banner_integrity,
            fonts_unchanged,
            style_ids_unchanged,
            media_unchanged,
        )
    )
    return {
        "passed": passed,
        "template_sha256": _sha256_bytes(template_docx),
        "generated_sha256": _sha256_bytes(generated_docx),
        "protected_part_hashes": protected_part_hashes,
        "protected_parts_present": protected_parts_present,
        "protected_parts_unchanged": protected_hashes_match,
        "header_parts_present": header_parts_present,
        "header_parts_unchanged": header_parts_unchanged,
        "header_rel_parts_present": header_rel_parts_present,
        "header_rel_parts_unchanged": header_rel_parts_unchanged,
        "banner_targets": banner_targets,
        "banner_integrity": banner_integrity,
        "fonts": {
            "template": template_fonts,
            "generated": generated_fonts,
            "unchanged": fonts_unchanged,
        },
        "style_ids": {
            "template_count": len(template_style_ids),
            "generated_count": len(generated_style_ids),
            "unchanged": style_ids_unchanged,
        },
        "media_unchanged": media_unchanged,
    }


def fill_template_docx(
    ts: TermSheet,
    template_path: str | Path = PREPROCESSED_TEMPLATE_PATH,
) -> tuple[bytes, dict[str, str]]:
    """Fill the pre-processed Paradigm template while preserving run formatting."""
    replacements = _build_replacements(ts)
    doc = Document(str(template_path))

    money_slots_used = 0
    percent_slots_used = 0

    for paragraph in _iter_all_paragraphs(doc):
        _replace_regex_all(paragraph, r"\[COMPANY\]", replacements["company_upper"])
        _replace_regex_all(
            paragraph,
            r"SERIES\s*\[\s*[_ ]\s*\]",
            f"SERIES {replacements['series']}",
        )
        _replace_regex_all(paragraph, r"Series\s*\[__\]", f"Series {replacements['series']}")
        _replace_regex_all(paragraph, r"\$\[1-10M\]M", replacements["debt_threshold"])
        _replace_regex_all(
            paragraph,
            r"\[Token Floor usually 50%\]",
            replacements["token_floor"],
        )

        _apply_co_investor_language(
            paragraph,
            ts.co_investor_language,
            ts.co_investor_text,
        )
        _apply_board_language(paragraph, ts.board_rights)

        while _replace_regex_once(
            paragraph,
            r"\$\[__\]M",
            (replacements["investment"] if money_slots_used == 0 else replacements["valuation"]),
        ):
            money_slots_used += 1

        while _replace_regex_once(
            paragraph,
            r"\[__\]\s*%",
            (
                f"{replacements['option_pool']}%"
                if percent_slots_used == 0
                else f"{replacements['ownership']}%"
            ),
        ):
            percent_slots_used += 1

        _apply_ipo_threshold(paragraph, replacements["ipo_threshold"])
        _apply_legal_fee_cap(paragraph, replacements["legal_fee_cap"])
        _apply_nvca_year(paragraph, replacements["nvca_year"])
        _apply_exclusivity_days(paragraph, replacements["exclusivity_days"])
        _apply_founder_carveout(paragraph, replacements["founder_carveout"])
        _apply_pro_rata_toggle(paragraph, ts.pro_rata_rights)

        _apply_seed_language(paragraph, ts.is_seed)
        _apply_protective_v_override(paragraph, ts.protective_provision_v_text)
        _apply_row_override(
            paragraph,
            "Customary NVCA investor rights, including information rights and pro rata rights",
            ts.other_rights_text,
        )
        _apply_row_override(
            paragraph,
            "For any Tokens (other than non-fungible tokens",
            ts.token_rights_text,
        )
        _apply_row_override(
            paragraph,
            "Founder vesting subject to due diligence.",
            ts.vesting_text,
        )

    if not ts.token_rights.enabled:
        _remove_token_rights_rows(doc)

    if money_slots_used != 2:
        raise ValueError(f"Expected 2 money slots in template, found {money_slots_used}")
    if percent_slots_used != 2:
        raise ValueError(f"Expected 2 percent slots in template, found {percent_slots_used}")

    buf = io.BytesIO()
    doc.save(buf)
    output = buf.getvalue()
    _assert_no_unresolved_markers(output)
    return output, replacements


def _docx_paragraph_texts(docx_bytes: bytes) -> list[str]:
    doc = Document(io.BytesIO(docx_bytes))
    return [p.text.strip() for p in _iter_all_paragraphs(doc) if p.text.strip()]


def _replacement_candidates(old_text: str, new_text: str) -> list[tuple[str, str]]:
    old_tokens = old_text.split()
    new_tokens = new_text.split()
    matcher = SequenceMatcher(a=old_tokens, b=new_tokens)
    candidates: list[tuple[str, str]] = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        old_frag = " ".join(old_tokens[i1:i2]).strip()
        new_frag = " ".join(new_tokens[j1:j2]).strip()
        if tag in {"replace", "delete"} and old_frag:
            candidates.append((old_frag, new_frag))
        if tag == "insert" and new_frag:
            prefix = " ".join(old_tokens[max(0, i1 - 6) : i1]).strip()
            suffix = " ".join(old_tokens[i1 : i1 + 6]).strip()
            anchor_old = " ".join(p for p in (prefix, suffix) if p).strip()
            anchor_new = " ".join(p for p in (prefix, new_frag, suffix) if p).strip()
            if anchor_old and anchor_new and anchor_old != anchor_new:
                candidates.append((anchor_old, anchor_new))
    return candidates


def generate_redline_docx(previous_docx: bytes, current_docx: bytes) -> bytes:
    previous_lines = _docx_paragraph_texts(previous_docx)
    current_lines = _docx_paragraph_texts(current_docx)
    matcher = SequenceMatcher(a=previous_lines, b=current_lines)

    replacements: list[tuple[str, str]] = []
    insertions: list[str] = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        if tag == "replace":
            old_block = previous_lines[i1:i2]
            new_block = current_lines[j1:j2]
            for idx in range(max(len(old_block), len(new_block))):
                old_t = old_block[idx] if idx < len(old_block) else ""
                new_t = new_block[idx] if idx < len(new_block) else ""
                if old_t and new_t and old_t != new_t:
                    replacements.append((old_t, new_t))
                elif old_t and not new_t:
                    replacements.append((old_t, ""))
                elif new_t and not old_t:
                    insertions.append(new_t)
        elif tag == "delete":
            for old_t in previous_lines[i1:i2]:
                replacements.append((old_t, ""))
        elif tag == "insert":
            insertions.extend(current_lines[j1:j2])

    replacements.sort(key=lambda item: len(item[0]), reverse=True)

    revision_doc = RevisionDocument(io.BytesIO(previous_docx))
    for old_text, new_text in replacements:
        if old_text == new_text:
            continue
        replaced = revision_doc.find_and_replace_tracked(old_text, new_text, author=_REDLINE_AUTHOR)
        if replaced > 0:
            continue
        for old_frag, new_frag in _replacement_candidates(old_text, new_text):
            if not old_frag or old_frag == new_frag:
                continue
            revision_doc.find_and_replace_tracked(old_frag, new_frag, author=_REDLINE_AUTHOR)

    if insertions and previous_lines:
        anchor = previous_lines[-1]
        appended = f"{anchor}\n" + "\n".join(insertions)
        revision_doc.find_and_replace_tracked(anchor, appended, author=_REDLINE_AUTHOR)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        temp_path = Path(handle.name)
    try:
        revision_doc.save(temp_path)
        return temp_path.read_bytes()
    finally:
        temp_path.unlink(missing_ok=True)


def convert_docx_to_pdf(docx_bytes: bytes, gotenberg_url: str | None = None) -> bytes:
    url = (gotenberg_url or os.getenv("GOTENBERG_URL") or "http://gotenberg:3000").rstrip("/")
    endpoint = f"{url}/forms/libreoffice/convert"
    with httpx.Client(timeout=60.0) as client:
        response = client.post(
            endpoint,
            files={
                "files": (
                    "document.docx",
                    docx_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        response.raise_for_status()
        return response.content


def generate_term_sheet_package(
    ts: TermSheet,
    *,
    previous_docx: bytes | None = None,
    include_pdf: bool = True,
    template_path: str | Path | None = None,
) -> GeneratedPackage:
    if ts.instrument_type != InstrumentType.PRICED:
        raise ValueError("DOCX package generation currently supports priced rounds only")
    resolved_template = Path(template_path) if template_path else TEMPLATE_PATH
    preprocessed = _ensure_preprocessed_template(resolved_template)

    clean_docx, replacements = fill_template_docx(ts, template_path=preprocessed)
    template_bytes = resolved_template.read_bytes()
    clean_docx = _restore_template_immutable_parts(template_bytes, clean_docx)
    _assert_template_structure_preserved(template_bytes, clean_docx)
    fidelity_report = _build_template_fidelity_report(template_bytes, clean_docx)
    redline_docx = generate_redline_docx(previous_docx, clean_docx) if previous_docx else None

    clean_pdf: bytes | None = None
    redline_pdf: bytes | None = None
    if include_pdf:
        try:
            clean_pdf = convert_docx_to_pdf(clean_docx)
            if redline_docx is not None:
                redline_pdf = convert_docx_to_pdf(redline_docx)
        except Exception:
            log.warning("PDF conversion unavailable — skipping", exc_info=True)

    return GeneratedPackage(
        clean_docx=clean_docx,
        redline_docx=redline_docx,
        clean_pdf=clean_pdf,
        redline_pdf=redline_pdf,
        replacements=replacements,
        fidelity_report=fidelity_report,
    )
